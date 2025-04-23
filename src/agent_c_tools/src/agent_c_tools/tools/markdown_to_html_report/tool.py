import json
import logging
import re

from pathlib import Path
from typing import Dict, List, Any

from agent_c.toolsets.tool_set import Toolset
from agent_c.toolsets.json_schema import json_schema

# Import dependencies for markdown to docx conversion
try:
    import markdown
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from bs4 import BeautifulSoup
    DOCX_CONVERSION_AVAILABLE = True
except ImportError:
    DOCX_CONVERSION_AVAILABLE = False
    logging.warning("python-docx and/or BeautifulSoup not available. The markdown to Word conversion will not work."
                    "Install with: pip install markdown python-docx beautifulsoup4")

logger = logging.getLogger(__name__)

# Regular expression to match UNC paths
UNC_PATH_PATTERN = r'^(//|\\\\)[^/\\]+(/|\\)[^/\\].*$'


class MarkdownToHtmlReportTools(Toolset):
    """Toolset for generating interactive HTML viewers from markdown files in a workspace."""

    def __init__(self, **kwargs):
        super().__init__(**kwargs, name="markdown_viewer", use_prefix=False)
        # Get workspace tools for file operations
        self.workspace_tool = self.tool_chest.active_tools.get('workspace')
        if not self.workspace_tool:
            logger.warning("Workspace toolset not available. This tool requires workspace tools.")

    @json_schema(
        description="Generate an interactive HTML viewer for markdown files in a workspace directory",
        params={
            "workspace": {
                "type": "string",
                "description": "The workspace containing the markdown files",
                "required": True
            },
            "input_path": {
                "type": "string",
                "description": "The path within the workspace where markdown files are located (or a full UNC path)",
                "required": True
            },
            "output_filename": {
                "type": "string",
                "description": "The name of the output HTML file to generate (can be a simple filename or full UNC path)",
                "required": True
            },
            "title": {
                "type": "string",
                "description": "Optional title for the HTML viewer (displayed in the sidebar)",
                "required": False
            },
            "files_to_ignore": {
                "type": "array",
                "items": {
                    "type": "string"
                },
                "description": "Optional flat list of filenames to ignore when generating the HTML viewer. Does not support folder level differentiation.",
                "required": False,
                "default": []
            }
        }
    )
    async def generate_md_viewer(self, **kwargs) -> str:
        """Generate an interactive HTML viewer for markdown files in a workspace directory.

        Args:
            workspace: The workspace containing the markdown files
            input_path: The path within the workspace where markdown files are located (or a full UNC path)
            output_filename: The name of the output HTML file to generate (can be a simple filename or full UNC path)
            title: Optional title for the HTML viewer

        Returns:
            A dictionary with success status and information about the operation
        """
        workspace = kwargs.get('workspace')
        input_path = kwargs.get('input_path')
        output_filename = kwargs.get('output_filename')
        title = kwargs.get('title', 'output_report.html')
        files_to_ignore = kwargs.get('files_to_ignore')

        # Arg check
        missing_fields = []
        if workspace is None:
            missing_fields.append("workspace")
        if input_path is None:
            missing_fields.append("input_path")
        if output_filename is None:
            missing_fields.append("output_filename")

        if missing_fields:
            return f"Required fields cannot be empty: {', '.join(missing_fields)}"

        combined_ignore_list = (files_to_ignore or []) + ([output_filename] if output_filename else [])

        try:
            # Check if input_path is already a UNC path
            if re.match(UNC_PATH_PATTERN, input_path):
                # Input path is already a UNC path, use it directly
                input_path_full = input_path
                # Extract workspace from UNC path if possible (for later use)
                try:
                    unc_parts = input_path.split('/', 3) if '/' in input_path else input_path.split('\\', 3)
                    if len(unc_parts) >= 3:
                        # Use the extracted workspace name
                        workspace = unc_parts[2]
                except Exception as e:
                    logger.warning(f"Could not extract workspace from UNC path: {e}")
            else:
                # If input_path is just the workspace name, don't append it again
                if input_path.strip('/') == workspace:
                    input_path_full = f"//{workspace}"
                else:
                    # Normalize input path for workspace construction
                    normalized_input_path = input_path.lstrip('/')
                    # Construct workspace UNC path
                    input_path_full = f"//{workspace}/{normalized_input_path}"

            # Ensure the output filename has .html extension
            if not output_filename.lower().endswith('.html'):
                output_filename += '.html'

            # Check if output_filename is already a UNC path
            if re.match(UNC_PATH_PATTERN, output_filename):
                # Output path is already a UNC path, use it directly
                output_path_full = output_filename
            else:
                # Normalize the output filename to remove any path information
                output_filename = Path(output_filename).name
                # Construct the output path in the workspace root
                output_path_full = f"//{workspace}/{output_filename}"

            logger.debug(f"Original input path: {input_path}")
            logger.debug(f"Workspace name: {workspace}")
            logger.debug(f"Processed input path: {input_path_full}")
            logger.debug(f"Output path: {output_path_full}")

            # Check if input directory exists by listing its contents
            logger.debug(f"Checking if input path exists...")
            ls_result = await self.workspace_tool.ls(path=input_path_full)
            ls_data = json.loads(ls_result)

            if 'error' in ls_data:
                return json.dumps({
                    "success": False,
                    "error": f"Input path '{input_path}' is not accessible: {ls_data['error']}"
                })

            # Start generating the viewer
            logger.debug(f"Collecting markdown files...")

            # Collect markdown files
            markdown_files = await self._collect_markdown_files(root_path=input_path_full,
                                                                files_to_ignore=combined_ignore_list)

            if not markdown_files:
                logger.debug(f"No markdown files found in {input_path}. Will search subdirectories recursively...")

                # Try to search deeper by examining subdirectories manually
                for item in ls_data.get('items', []):
                    if item['type'] == 'directory':
                        subdir_path = f"{input_path_full}/{item['name']}"
                        subdir_files = await self._collect_markdown_files(subdir_path)
                        if subdir_files:
                            markdown_files.update(subdir_files)

                # If still no files found
                if not markdown_files:
                    return json.dumps({
                        "success": False,
                        "error": f"No markdown files found in {input_path} or its subdirectories"
                    })

            logger.debug(f"Found {len(markdown_files)} markdown files. Building file structure...")

            # Build file structure for the viewer
            file_structure = await self._build_file_structure(input_path_full, markdown_files)

            if not file_structure:
                return json.dumps({
                    "success": False,
                    "error": "Failed to build file structure"
                })

            # Get the HTML template
            logger.debug("Preparing HTML template...")
            html_template = await self._get_html_template()

            # Customize title if provided
            if title:
                html_template = html_template.replace(
                    "<h3 style=\"margin: 0 16px 16px 16px;\">Agent C Output Viewer</h3>",
                    f"<h3 style=\"margin: 0 16px 16px 16px;\">{title}</h3>")

            # Replace placeholder with file structure
            json_structure = json.dumps(file_structure)
            html_content = html_template.replace('$FILE_STRUCTURE', json_structure)

            # Write the generated HTML to the workspace
            logger.debug(f"Writing HTML viewer to output location...")
            write_result = await self.workspace_tool.write(
                path=output_path_full,
                data=html_content,
                mode="write"
            )

            write_data = json.loads(write_result)
            if 'error' in write_data:
                return json.dumps({
                    "success": False,
                    "error": f"Failed to write HTML file: {write_data['error']}"
                })

            message = f"Successfully generated HTML viewer at {output_filename}."
            logger.debug(message)
            
            # Raise a media event with HTML that provides a link to the generated file
            try:
                # Create a fully resolved file path for browser access
                # This will be a local file URL that should open in the browser
                file_url = f"file://{output_path_full.replace('//','/')}"

                # Get the actual OS-level filepath using the workspace's full_path method

                _, workspace_obj, rel_path = self.workspace_tool._parse_unc_path(output_path_full)
                file_system_path = None
                if workspace_obj and hasattr(workspace_obj, 'full_path'):
                    # The full_path method handles path normalization and joining with workspace root
                    # Set mkdirs=False since we're just getting the path for a URL, not writing
                    file_system_path = workspace_obj.full_path(rel_path, mkdirs=False)

                # Create a file:// URL from the system path
                if file_system_path:
                    # Convert backslashes to forward slashes for URL
                    url_path = file_system_path.replace('\\', '/')

                    # Ensure correct URL format (need 3 slashes for file:// URLs with absolute paths)
                    if url_path.startswith('/'):
                        file_url = f"file://{url_path}"
                    else:
                        file_url = f"file:///{url_path}"
                else:
                    # Fallback if we couldn't get the actual path
                    file_url = f"file:///{output_path_full.replace('//', '').replace('\\', '/')}"

                # Create HTML content with a link that opens in a new window
                # html_content = f"<a href='{file_url}' target='_blank'>Open Report: {output_filename}</a>"

                html_content = f"""
                <div style="padding: 20px; font-family: system-ui, sans-serif; max-width: 800px; margin: 0 auto; border: 1px solid #e2e8f0; border-radius: 8px; background-color: #f8fafc;">
                    <h2 style="color: #334155; margin-top: 0;">Report Generated Successfully</h2>

                    <div style="background-color: #f1f5f9; border-radius: 6px; padding: 16px; margin-bottom: 20px;">
                        <p style="margin: 0 0 8px 0;"><strong>File:</strong> {output_filename}</p>
                        <p style="margin: 0 0 8px 0;"><strong>Contents:</strong> {len(markdown_files)} markdown files</p>
                        <p style="margin: 0;"><strong>Location:</strong> <code style="background: #e2e8f0; padding: 2px 4px; border-radius: 4px;">{output_path_full}</code></p>
                    </div>

                    <div style="background-color: #fff7ed; border-left: 4px solid #f97316; padding: 16px; margin-bottom: 20px;">
                        <p style="margin: 0; font-weight: 500; color: #9a3412;">Browser Security Notice</p>
                        <p style="margin: 8px 0 0 0;">Due to browser security restrictions, you'll need to manually open the file:</p>
                    </div>

                    <div style="margin-bottom: 16px;">
                        <p><strong>File path:</strong> <br/>
                        <code style="background: #e2e8f0; padding: 8px; border-radius: 4px; display: block; margin-top: 8px; word-break: break-all;">{file_system_path}</code>
                        </p>

                        <p><strong>Terminal command:</strong> <br/>
                        <code style="background: #e2e8f0; padding: 8px; border-radius: 4px; display: block; margin-top: 8px; word-break: break-all;">
                            {f'start "" "{file_system_path}"' if ':\\' in file_system_path else f'open "{file_system_path}"'}
                        </code>
                        </p>
                    </div>

                    <ol style="margin-top: 0; padding-left: 20px;">
                        <li style="margin-bottom: 8px;">Copy the file path and paste it into your file explorer address bar</li>
                        <li style="margin-bottom: 8px;">Or copy the command and run it in your terminal/command prompt</li>
                    </ol>

                    <div style="margin-top: 16px;">
                        <a href="{file_url}" target="_blank" style="display: inline-block; background-color: #3b82f6; color: white; text-decoration: none; padding: 10px 16px; border-radius: 6px; font-weight: 500;">Try Direct Link</a>
                        <span style="margin-left: 8px; color: #6b7280;">(may not work due to browser restrictions)</span>
                    </div>
                </div>
                """
                # Raise the media event
                await self._raise_render_media(
                    sent_by_class=self.__class__.__name__,
                    sent_by_function='generate_report',
                    content_type="text/html",
                    content=html_content
                )

                # markdown_content = f"""# Your Report is Ready
                # The report containing **{len(markdown_files)} markdown files** has been generated.
                # ## Details
                # - Filename: `{output_filename}`
                # - Location: `{file_url}`
                #
                # [Click here to open the report]({file_url})"""
                #
                # await self._raise_render_media(
                #     sent_by_class=self.__class__.__name__,
                #     sent_by_function='generate_report',
                #     content_type="text/markdown",  # Use text/markdown content type
                #     content=markdown_content
                # )
                
                # logger.debug(f"Raised media event with link to generated HTML file: {file_url}")
            except Exception as e:
                logger.error(f"Failed to raise media event: {str(e)}")

            return json.dumps({
                "success": True,
                "message": message,
                "output_file": output_filename,
                "output_path": output_path_full,
                "workspace": workspace,
                "file_count": len(markdown_files)
            })

        except Exception as e:
            logger.exception("Error generating markdown viewer")
            return json.dumps({
                "success": False,
                "error": f"Error generating markdown viewer: {str(e)}"
            })

    async def _collect_markdown_files(self, root_path: str, files_to_ignore: List[str] = None) -> Dict[str, str]:
        """Collect all markdown files in the workspace directory structure."""
        markdown_files = {}
        files_to_ignore = [f.lower() for f in (files_to_ignore or [])]
        logger.debug(f"Searching for markdown files in {root_path} and subdirectories...")

        # Process a directory to find markdown files recursively
        async def process_directory(dir_path, rel_path=""):
            # Ensure proper UNC path format for directory listing
            normalized_dir_path = dir_path.replace('\\', '/')
            # Ensure trailing slash for consistent path handling
            if not normalized_dir_path.endswith('/'):
                normalized_dir_path += '/'

            try:
                ls_result = await self.workspace_tool.ls(path=normalized_dir_path)
                ls_data = json.loads(ls_result)

                if 'error' in ls_data:
                    logger.error(f"Error listing directory '{normalized_dir_path}': {ls_data['error']}")
                    return

                items = ls_data.get('contents', [])
                logger.debug(f"Found {len(items)} items in {normalized_dir_path}")

                for item_name in items:
                    # Skip hidden files and directories
                    if item_name.startswith('.'):
                        continue

                    if item_name.lower() in files_to_ignore:
                        continue

                    # Create proper UNC item path - don't add extra slash if dir already ends with /
                    item_path = f"{normalized_dir_path}{item_name}"

                    # First check if it's a markdown file by extension - only process .md and .markdown files
                    if self._is_markdown_file(item_name):
                        # Add markdown file to our collection
                        file_rel_path = f"{rel_path}/{item_name}" if rel_path else item_name
                        # Normalize path separators in keys
                        normalized_file_rel_path = file_rel_path.replace('\\', '/')
                        markdown_files[normalized_file_rel_path] = item_path
                        # logger.debug(f"Found markdown file: {normalized_file_rel_path}")
                    else:
                        # Only check if it's a directory - we don't process non-markdown files
                        is_dir_result = await self.workspace_tool.is_directory(path=item_path)
                        is_dir_data = json.loads(is_dir_result)

                        if 'error' not in is_dir_data and is_dir_data.get('is_directory', False):
                            # Recursively process subdirectory
                            new_rel_path = f"{rel_path}/{item_name}" if rel_path else item_name
                            await process_directory(item_path, new_rel_path)
                        # No else clause here - we don't add non-markdown files
            except Exception as e:
                logger.error(f"Error processing directory {normalized_dir_path}: {e}")

        # Start processing from the root path, ensuring proper UNC format
        normalized_root_path = root_path.replace('\\', '/')
        # Make sure root path doesn't have a trailing slash for initial call
        normalized_root_path = normalized_root_path.rstrip('/')
        await process_directory(normalized_root_path)

        file_count = len(markdown_files)
        logger.info(f"Found {file_count} markdown files in {normalized_root_path}")
        logger.debug(f"Found {file_count} markdown files in total")
        return markdown_files

    def _is_markdown_file(self, filename: str) -> bool:
        """Check if a file is a markdown file."""
        filename_lower = filename.lower()
        return filename_lower.endswith('.md') or filename_lower.endswith('.markdown')

    async def _build_file_structure(self, root_path: str, markdown_files: Dict[str, str]) -> List[Dict[str, Any]]:
        """Build file structure for the HTML viewer."""
        structure = []
        folders = {}

        # Normalize root path for consistent path handling
        normalized_root_path = root_path.replace('\\', '/')

        # Sort paths to ensure proper order (parent folders before children)
        for rel_path, unc_path in sorted(markdown_files.items()):
            # Double-check that it's a markdown file before processing
            if not self._is_markdown_file(Path(rel_path).name):
                logger.debug(f"Skipping non-markdown file: {rel_path}")
                continue

            # Normalize the UNC path for consistent handling
            normalized_unc_path = unc_path.replace('\\', '/')

            # Get path parts for folder hierarchy construction
            path_parts = rel_path.split('/')
            file_name = path_parts[-1]

            # Create folder hierarchy
            current_folders = structure
            current_path = ""

            # Process all path parts except the last one (which is the file)
            for part in path_parts[:-1]:
                if current_path:
                    current_path += f"/{part}"
                else:
                    current_path = part

                # Check if this folder already exists in our structure
                folder = next((f for f in current_folders if f['type'] == 'folder' and f['name'] == part), None)

                if not folder:
                    # Create new folder
                    folder = {
                        'name': part,
                        'type': 'folder',
                        'path': current_path,
                        'children': []
                    }
                    current_folders.append(folder)
                    folders[current_path] = folder['children']

                # Move to the next level
                current_folders = folder['children']

            try:
                # Read the file content
                read_result = await self.workspace_tool.read(path=normalized_unc_path)
                read_data = json.loads(read_result)

                if 'error' in read_data:
                    logger.error(f"Error reading file {normalized_unc_path}: {read_data['error']}")
                    file_content = f"Error reading file: {read_data['error']}"
                else:
                    file_content = read_data.get('contents', '')

                # Process links in the content
                processed_content = await self._process_markdown_links(file_content, rel_path, normalized_root_path,
                                                                       markdown_files)

                # Add the file to the structure
                current_folders.append({
                    'name': file_name,
                    'type': 'file',
                    'path': rel_path,
                    'content': processed_content
                })
            except Exception as e:
                logger.error(f"Error processing file {normalized_unc_path}: {e}")
                # Add the file with error message
                current_folders.append({
                    'name': file_name,
                    'type': 'file',
                    'path': rel_path,
                    'content': f"Error processing file: {str(e)}"
                })

        return structure

    async def _process_markdown_links(self, content: str, file_rel_path: str, root_path: str,
                                      markdown_files: Dict[str, str]) -> str:
        """Convert markdown links to other markdown files into internal viewer links."""
        # Normalize path separators in file_rel_path
        normalized_file_rel_path = file_rel_path.replace('\\', '/')

        # Get the directory of the current file for resolving relative paths
        file_dir = Path(normalized_file_rel_path).parent.as_posix() if '/' in normalized_file_rel_path else ""

        # Regular expression to find markdown links
        link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'

        def replace_link(match):
            link_text = match.group(1)
            link_target = match.group(2)

            # Handle internal document links (just anchors)
            if link_target.startswith('#'):
                anchor_id = link_target.replace("#", "")
                return f'[{link_text}](javascript:void(scrollToAnchor("{anchor_id}")))'

            # Skip already processed javascript links
            if link_target.startswith('javascript:'):
                return match.group(0)

            # Skip external links
            if link_target.startswith(('http://', 'https://')):
                return match.group(0)

            # Split target into file path and anchor if it exists
            parts = link_target.split('#', 1)
            target_file = parts[0]
            anchor = f"#{parts[1]}" if len(parts) > 1 else ""

            try:
                # Normalize target file path separators
                target_file = target_file.replace('\\', '/')

                # Handle UNC paths in links
                if re.match(UNC_PATH_PATTERN, target_file):
                    # If the link target is a UNC path, extract the relative part
                    try:
                        # Extract the part after the workspace
                        unc_parts = target_file.split('/', 3)
                        if len(unc_parts) >= 4:
                            target_path = unc_parts[3]
                        else:
                            # Can't extract relative path, use as is
                            target_path = target_file
                    except Exception:
                        target_path = target_file
                else:
                    # Calculate the full path from the current file's directory
                    if target_file.startswith('/'):
                        # Handle absolute path (relative to root)
                        target_path = target_file[1:]
                    else:
                        # Handle relative path
                        if file_dir:
                            target_path = f"{file_dir}/{target_file}"
                        else:
                            target_path = target_file

                # Normalize path and separators
                target_path = Path(target_path).as_posix()

                # Check if the target is a markdown file in our collection
                if target_path in markdown_files:
                    # Replace with internal link
                    if anchor:
                        return f'[{link_text}](javascript:void(openMarkdownFile("{target_path}", "{anchor}")))'
                    else:
                        return f'[{link_text}](javascript:void(openMarkdownFile("{target_path}")))'

                # If it's potentially a directory link (ending with / or no extension)
                if target_path.endswith('/') or '.' not in target_path.split('/')[-1]:
                    # Ensure path ends with a slash for proper directory prefix matching
                    if not target_path.endswith('/'):
                        target_path = f"{target_path}/"

                    # Check if there are any markdown files in this directory
                    for md_path in markdown_files.keys():
                        if md_path.startswith(target_path):
                            # It's a directory with markdown files, link to the first one
                            return f'[{link_text}](javascript:void(openMarkdownFile("{md_path}")))'
            except Exception as e:
                # If there's an error resolving the path, log it and leave the link unchanged
                logger.error(f"Error processing link '{link_target}' in file '{normalized_file_rel_path}': {e}")

            # If target doesn't exist in our structure or path resolution failed, leave the link unchanged
            return match.group(0)

        # Process all links in the content
        processed_content = re.sub(link_pattern, replace_link, content)
        return processed_content

    async def _get_html_template(self) -> str:
        """Get the HTML template for the markdown viewer."""
        try:
            # If no workspace template found, try local file system
            local_template_path = Path(__file__).parent / "markdown-viewer-template.html"

            # Check if the template file exists
            if not local_template_path.exists():
                logger.warning(f"Template file not found: {local_template_path}")
                # Try alternate locations in case __file__ is not working as expected
                alt_paths = [
                    Path.cwd() / "markdown-viewer-template.html",
                    Path.cwd() / "tools" / "markdown_viewer" / "markdown-viewer-template.html",
                ]

                for alt_path in alt_paths:
                    if alt_path.exists():
                        local_template_path = alt_path
                        logger.info(f"Found template at alternate location: {local_template_path}")
                        break
                else:
                    raise FileNotFoundError("Could not find template file in workspace or local paths")

            # Read the template file from local file system
            with open(local_template_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading template file: {e}")
            raise RuntimeError(f"Failed to load HTML template: {e}")


    @json_schema(
        description="Convert a markdown file to Word (DOCX) format",
        params={
            "workspace": {
                "type": "string",
                "description": "The workspace containing the markdown file",
                "required": True
            },
            "input_path": {
                "type": "string",
                "description": "The path to the markdown file to convert (UNC-style or relative to workspace)",
                "required": True
            },
            "output_filename": {
                "type": "string",
                "description": "filename for the output Word document",
                "required": False
            },
            "style": {
                "type": "string",
                "description": "Style template to use",
                "enum": ["default", "academic", "business", "minimal"],
                "required": False,
                "default": "default"
            },
            "include_toc": {
                "type": "boolean",
                "description": "Whether to include a table of contents",
                "required": False,
                "default": True
            },
            "page_break_level": {
                "type": "integer",
                "description": "Insert page breaks before headings of this level (1-6)",
                "required": False,
                "default": 1
            }
        }
    )
    async def markdown_to_docx(self, **kwargs) -> str:
        """Convert a markdown file to Word (DOCX) format.

        Args:
            kwargs:
                workspace: The workspace containing the markdown file
                input_path: The path to the markdown file to convert (UNC-style or relative to workspace)
                output_path: Output path for the Word document (UNC-style or relative to workspace)
                style: Style template to use (default, academic, business, minimal)
                include_toc: Whether to include a table of contents
                page_break_level: Insert page breaks before headings of this level (1-6)

        Returns:
            A dictionary with success status and information about the operation
        """
        if not DOCX_CONVERSION_AVAILABLE:
            return json.dumps({
                "success": False,
                "error": "Required dependencies not available. Please install python-markdown, python-docx, and beautifulsoup4."
            })

        workspace = kwargs.get('workspace')
        input_path = kwargs.get('input_path')
        output_filename = kwargs.get('output_filename')
        style = kwargs.get('style', 'default')
        include_toc = kwargs.get('include_toc', True)
        page_break_level = kwargs.get('page_break_level', 1)

        # Arg check
        missing_fields = []
        if workspace is None:
            missing_fields.append("workspace")
        if input_path is None:
            missing_fields.append("input_path")
        if output_filename is None:
            missing_fields.append("output_filename")

        if missing_fields:
            return f"Required fields cannot be empty: {', '.join(missing_fields)}"

        try:
            # Check if input_path is already a UNC path (starts with //)
            if input_path.startswith('//'):
                # Input path is already a UNC path, use it directly
                input_path_full = input_path

                # Extract workspace from UNC path
                path_parts = input_path.replace('\\', '/').split('/', 4)
                if len(path_parts) >= 3:
                    workspace = path_parts[2]  # Get workspace name
            else:
                # Normalize input path (remove leading slashes)
                normalized_input_path = input_path.lstrip('/')

                # Construct full UNC path
                input_path_full = f"//{workspace}/{normalized_input_path}"

            # Extract the filename using string manipulation
            # Convert all backslashes to forward slashes for consistent handling
            normalized_path = input_path_full.replace('\\', '/')
            # Get the last part after the final slash
            input_filename = normalized_path.split('/')[-1]

            if not self._is_markdown_file(normalized_path):
                return json.dumps({
                    "success": False,
                    "error": f"Input file '{normalized_path}' does not appear to be a markdown file. Expected a .md or .markdown extension."
                })

            # Determine output path if not provided
            if not output_filename:
                # Use the same name as the input but with .docx extension
                output_filename = Path(input_filename).stem + ".docx"
                output_path_full = f"//{workspace}/{output_filename}"
            else:
                # Check if output_path is already a UNC path
                if re.match(UNC_PATH_PATTERN, output_filename):
                    output_path_full = output_filename
                else:
                    # Normalize the output path
                    normalized_output_path = output_filename.lstrip('/')
                    # Ensure it has .docx extension
                    if not normalized_output_path.lower().endswith('.docx'):
                        normalized_output_path = Path(normalized_output_path).with_suffix('.docx').as_posix()
                    output_path_full = f"//{workspace}/{normalized_output_path}"

            # Log the paths for debugging
            logger.debug(f"Original input path: {input_path}")
            logger.debug(f"Workspace name: {workspace}")
            logger.debug(f"Processed input path: {input_path_full}")
            logger.debug(f"Output path: {output_path_full}")

            # Read and Process the markdown file
            read_result = await self.workspace_tool.read(path=input_path_full)
            read_data = json.loads(read_result)

            if 'error' in read_data:
                return json.dumps({
                    "success": False,
                    "error": f"Failed to read input file: {read_data['error']}"
                })

            md_content = read_data.get('contents', '')

            # Convert markdown to Word document
            docx_content_bytes = await self._convert_markdown_to_docx_bytes(md_content, style, include_toc, page_break_level)

            try:
                write_result = await self.workspace_tool.internal_write_bytes(
                    path=output_path_full,
                    data=docx_content_bytes,
                    mode="write",
                )
            except Exception as e:
                logger.error(f"Error writing docx file: {e}")
                return json.dumps({
                    "success": False,
                    "error": f"Failed to write Word document: {str(e)}"
                })

            write_data = json.loads(write_result)
            if 'error' in write_data:
                return json.dumps({
                    "success": False,
                    "error": f"Failed to write Word document: {write_data['error']}"
                })

            # Prepare output summary
            result = {
                "success": True,
                "message": f"Successfully converted markdown to Word document at {output_path_full}",
                "input_file": input_path_full,
                "output_file": output_path_full,
                "workspace": workspace,
                "style": style
            }

            # Raise a media event with HTML that provides information about the conversion
            try:
                # Get the actual OS-level filepath for the output
                _, workspace_obj, rel_path = self.workspace_tool._parse_unc_path(output_path_full)
                file_system_path = None
                if workspace_obj and hasattr(workspace_obj, 'full_path'):
                    file_system_path = workspace_obj.full_path(rel_path, mkdirs=False)

                # Create file URL (this may not work due to browser security)
                if file_system_path:
                    url_path = file_system_path.replace('\\', '/')
                    if url_path.startswith('/'):
                        file_url = f"file://{url_path}"
                    else:
                        file_url = f"file:///{url_path}"
                else:
                    file_url = f"file:///{output_path_full.replace('//', '').replace('\\', '/')}"

                html_content = f"""
                <div style="padding: 20px; font-family: system-ui, sans-serif; max-width: 800px; margin: 0 auto; border: 1px solid #e2e8f0; border-radius: 8px; background-color: #f8fafc;">
                    <h2 style="color: #334155; margin-top: 0;">Word Document Generated Successfully</h2>

                    <div style="background-color: #f1f5f9; border-radius: 6px; padding: 16px; margin-bottom: 20px;">
                        <p style="margin: 0 0 8px 0;"><strong>Input:</strong> {input_path_full}</p>
                        <p style="margin: 0 0 8px 0;"><strong>Output:</strong> {output_path_full}</p>
                        <p style="margin: 0 0 8px 0;"><strong>Style:</strong> {style}</p>
                        <p style="margin: 0;"><strong>Location:</strong> <code style="background: #e2e8f0; padding: 2px 4px; border-radius: 4px;">{file_system_path if file_system_path else output_path_full}</code></p>
                    </div>

                    <div style="margin-bottom: 16px;">
                        <p><strong>File path:</strong> <br/>
                        <code style="background: #e2e8f0; padding: 8px; border-radius: 4px; display: block; margin-top: 8px; word-break: break-all;">{file_system_path if file_system_path else output_path_full}</code>
                        </p>

                        <p><strong>Terminal command:</strong> <br/>
                        <code style="background: #e2e8f0; padding: 8px; border-radius: 4px; display: block; margin-top: 8px; word-break: break-all;">
                            {f'start "" "{file_system_path}"' if file_system_path and ':\\' in file_system_path else f'open "{file_system_path if file_system_path else output_path_full}"'}
                        </code>
                        </p>
                    </div>
                </div>
                """
                # Raise the media event
                await self._raise_render_media(
                    sent_by_class=self.__class__.__name__,
                    sent_by_function='markdown_to_docx',
                    content_type="text/html",
                    content=html_content
                )
                
            except Exception as e:
                logger.error(f"Failed to raise media event: {str(e)}")

            return json.dumps(result)

        except Exception as e:
            logger.exception("Error converting markdown to Word document")
            return json.dumps({
                "success": False,
                "error": f"Error converting markdown to Word document: {str(e)}"
            })

    async def _convert_markdown_to_docx_bytes(self, md_content: str, style_name: str = "default",
                                   include_toc: bool = True, page_break_level: int = 1) -> bytes:
        """Convert markdown content to a Word document in binary format.
        
        Note: This method returns binary data which will need to be encoded as base64
        when writing to the workspace since the workspace API primarily works with text.


        Args:
            md_content: The markdown content to convert
            style_name: Style template to use (default, academic, business, minimal)
            include_toc: Whether to include a table of contents
            page_break_level: Insert page breaks before headings of this level

        Returns:
            Bytes containing the Word document
        """
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_content,
            extensions=['extra', 'codehilite', 'toc', 'tables', 'nl2br', 'sane_lists']
        )

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')

        # Create new Word document
        doc = Document()
        doc = self._apply_style(doc, style_name)

        # Add table of contents if requested
        if include_toc:
            self._create_table_of_contents(doc)

        # Process HTML elements
        current_elements = soup.find_all(
            ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'pre', 'blockquote', 'table'])

        for element in current_elements:
            tag_name = element.name

            # Process headings
            if tag_name.startswith('h') and len(tag_name) == 2:
                level = int(tag_name[1])

                # Add page break before heading if configured
                if level <= page_break_level:
                    doc.add_page_break()

                # Add heading
                doc.add_heading(element.get_text(), level=level)

            # Process paragraphs
            elif tag_name == 'p':
                paragraph = doc.add_paragraph()

                # Process inline elements (bold, italic, links, etc.)
                for child in element.children:
                    if child.name is None:  # Plain text
                        paragraph.add_run(child.string)
                    elif child.name == 'strong' or child.name == 'b':
                        paragraph.add_run(child.get_text()).bold = True
                    elif child.name == 'em' or child.name == 'i':
                        paragraph.add_run(child.get_text()).italic = True
                    elif child.name == 'a':
                        self._add_hyperlink(paragraph, child.get('href', ''), child.get_text())
                    elif child.name == 'code':
                        run = paragraph.add_run(child.get_text())
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    elif child.name == 'br':
                        paragraph.add_run('\n')
                    else:
                        paragraph.add_run(child.get_text())

            # Process lists
            elif tag_name in ('ul', 'ol'):
                is_ordered = tag_name == 'ol'

                for list_item in element.find_all('li', recursive=False):
                    level = 0
                    parent = list_item.parent
                    while parent is not None and parent.name in ('ul', 'ol'):
                        level += 1
                        parent = parent.parent

                    # Add list item
                    p = doc.add_paragraph(style='List Bullet' if not is_ordered else 'List Number')
                    p.paragraph_format.left_indent = Inches(0.25 * level)

                    # Process inline elements in list items
                    item_text = ""
                    for child in list_item.children:
                        if child.name is None:  # Plain text
                            if child.string:
                                item_text += child.string.strip()
                        elif child.name == 'strong' or child.name == 'b':
                            p.add_run(child.get_text().strip()).bold = True
                        elif child.name == 'em' or child.name == 'i':
                            p.add_run(child.get_text().strip()).italic = True
                        elif child.name not in ('ul', 'ol'):  # Avoid adding nested list text twice
                            item_text += child.get_text().strip() + " "

                    if item_text:
                        p.add_run(item_text.strip())

                    # Handle nested lists recursively (simplified)
                    nested_lists = list_item.find_all(['ul', 'ol'], recursive=False)
                    for nested_list in nested_lists:
                        for nested_item in nested_list.find_all('li', recursive=False):
                            nested_p = doc.add_paragraph(style='List Bullet' if nested_list.name == 'ul' else 'List Number')
                            nested_p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                            nested_p.add_run(nested_item.get_text())

            # Process code blocks
            elif tag_name == 'pre':
                code_block = element.find('code')
                if code_block:
                    p = doc.add_paragraph()
                    code_text = code_block.get_text()

                    # Apply code block styling
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)

                    # Add a light gray shading
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)

            # Process blockquotes
            elif tag_name == 'blockquote':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                p.style = 'Quote'
                p.add_run(element.get_text())

            # Process tables
            elif tag_name == 'table':
                # Count rows and columns
                rows = element.find_all('tr')
                if not rows:
                    continue

                # Determine max columns by checking all rows
                max_cols = 0
                for row in rows:
                    cols = row.find_all(['td', 'th'])
                    max_cols = max(max_cols, len(cols))

                if max_cols == 0:
                    continue

                # Create table
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'

                # Fill table
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j < max_cols:
                            # Style header cells
                            if cell.name == 'th':
                                cell_text = cell.get_text().strip()
                                table.cell(i, j).text = cell_text
                                for paragraph in table.cell(i, j).paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                            else:
                                table.cell(i, j).text = cell.get_text().strip()

        # Save document to a BytesIO object
        from io import BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue()

    def _apply_style(self, doc: Document, style_name: str) -> Document:
        """Apply predefined style to document."""
        # Default style is already applied by default

        if style_name == 'academic':
            # Font settings
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal'].font.size = Pt(12)

            # Heading styles
            for i in range(1, 4):
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = 'Times New Roman'
                heading_style.font.bold = True
                heading_style.font.size = Pt(16 - (i - 1) * 2)  # H1: 16pt, H2: 14pt, H3: 12pt

            # Paragraph spacing
            doc.styles['Normal'].paragraph_format.line_spacing = 2.0  # Double spacing

        elif style_name == 'business':
            # Font settings
            doc.styles['Normal'].font.name = 'Calibri'
            doc.styles['Normal'].font.size = Pt(11)

            # Heading styles
            for i in range(1, 4):
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = 'Calibri'
                heading_style.font.bold = True
                if i == 1:
                    heading_style.font.size = Pt(16)
                    heading_style.font.color.rgb = RGBColor(0, 77, 113)  # Dark blue
                elif i == 2:
                    heading_style.font.size = Pt(14)
                    heading_style.font.color.rgb = RGBColor(0, 112, 155)  # Medium blue
                else:
                    heading_style.font.size = Pt(12)
                    heading_style.font.color.rgb = RGBColor(0, 130, 188)  # Light blue

        elif style_name == 'minimal':
            # Font settings
            doc.styles['Normal'].font.name = 'Arial'
            doc.styles['Normal'].font.size = Pt(10)

            # Heading styles
            for i in range(1, 4):
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = 'Arial'
                heading_style.font.bold = True
                heading_style.font.size = Pt(14 - (i - 1) * 2)  # H1: 14pt, H2: 12pt, H3: 10pt

            # Paragraph spacing
            doc.styles['Normal'].paragraph_format.space_after = Pt(6)

        return doc

    def _create_table_of_contents(self, doc: Document) -> None:
        """Create table of contents"""
        doc.add_heading('Table of Contents', level=1)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update table of contents."

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)
        r_element.append(fldChar4)

        doc.add_page_break()

    def _add_hyperlink(self, paragraph, url, text):
        """Add a hyperlink to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                              is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0000FF')
        rPr.append(c)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink


# Register the toolset with the Agent C framework
Toolset.register(MarkdownToHtmlReportTools)
