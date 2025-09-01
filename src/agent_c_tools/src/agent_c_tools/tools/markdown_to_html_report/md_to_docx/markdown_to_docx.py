import logging


logger = logging.getLogger(__name__)

try:
    import markdown
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from bs4 import BeautifulSoup
    DOCX_CONVERSION_AVAILABLE = True
except ImportError:
    DOCX_CONVERSION_AVAILABLE = False
    logging.warning("python-docx and/or BeautifulSoup not available. The markdown to Word conversion will not work."
                    "Install with: pip install markdown python-docx beautifulsoup4")

class MarkdownToDocxConverter:
    """Helper class for converting markdown to Word documents."""

    def __init__(self):
        if not DOCX_CONVERSION_AVAILABLE:
            self.docx_conversion_available = False
            logger.warning("Document conversion dependencies not available.")
        else:
            self.docx_conversion_available = True

    async def convert_to_docx(self, md_content: str, style_name: str = "default",
                              include_toc: bool = True, page_break_level: int = 1) -> bytes:
        """Convert markdown content to a Word document in binary format."""
        if not DOCX_CONVERSION_AVAILABLE:
            raise ImportError(
                "Required dependencies not available. Please install python-markdown, python-docx, and beautifulsoup4.")

        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_content,
            extensions=['extra', 'codehilite', 'toc', 'tables', 'nl2br', 'sane_lists']
        )

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')

        # Create new Word document
        doc = Document()
        doc = self._apply_style(doc, style_name)

        # Add table of contents if requested
        if include_toc:
            self._create_table_of_contents(doc)

        # Process HTML elements
        current_elements = soup.find_all(
            ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'pre', 'blockquote', 'table'])

        for element in current_elements:
            tag_name = element.name

            # Process headings
            if tag_name.startswith('h') and len(tag_name) == 2:
                level = int(tag_name[1])

                # Add page break before heading if configured
                if level <= page_break_level:
                    doc.add_page_break()

                # Add heading
                doc.add_heading(element.get_text(), level=level)

            # Process paragraphs
            elif tag_name == 'p':
                paragraph = doc.add_paragraph()

                # Process inline elements (bold, italic, links, etc.)
                for child in element.children:
                    if child.name is None:  # Plain text
                        paragraph.add_run(child.string)
                    elif child.name == 'strong' or child.name == 'b':
                        paragraph.add_run(child.get_text()).bold = True
                    elif child.name == 'em' or child.name == 'i':
                        paragraph.add_run(child.get_text()).italic = True
                    elif child.name == 'a':
                        self._add_hyperlink(paragraph, child.get('href', ''), child.get_text())
                    elif child.name == 'code':
                        run = paragraph.add_run(child.get_text())
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    elif child.name == 'br':
                        paragraph.add_run('\n')
                    else:
                        paragraph.add_run(child.get_text())

            # Process lists
            elif tag_name in ('ul', 'ol'):
                is_ordered = tag_name == 'ol'

                for list_item in element.find_all('li', recursive=False):
                    level = 0
                    parent = list_item.parent
                    while parent is not None and parent.name in ('ul', 'ol'):
                        level += 1
                        parent = parent.parent

                    # Add list item
                    p = doc.add_paragraph(style='List Bullet' if not is_ordered else 'List Number')
                    p.paragraph_format.left_indent = Inches(0.25 * level)

                    # Process inline elements in list items
                    item_text = ""
                    for child in list_item.children:
                        if child.name is None:  # Plain text
                            if child.string:
                                item_text += child.string.strip()
                        elif child.name == 'strong' or child.name == 'b':
                            p.add_run(child.get_text().strip()).bold = True
                        elif child.name == 'em' or child.name == 'i':
                            p.add_run(child.get_text().strip()).italic = True
                        elif child.name not in ('ul', 'ol'):  # Avoid adding nested list text twice
                            item_text += child.get_text().strip() + " "

                    if item_text:
                        p.add_run(item_text.strip())

                    # Handle nested lists recursively (simplified)
                    nested_lists = list_item.find_all(['ul', 'ol'], recursive=False)
                    for nested_list in nested_lists:
                        for nested_item in nested_list.find_all('li', recursive=False):
                            nested_p = doc.add_paragraph(
                                style='List Bullet' if nested_list.name == 'ul' else 'List Number')
                            nested_p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                            nested_p.add_run(nested_item.get_text())

            # Process code blocks
            elif tag_name == 'pre':
                code_block = element.find('code')
                if code_block:
                    p = doc.add_paragraph()
                    code_text = code_block.get_text()

                    # Apply code block styling
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)

                    # Add a light gray shading
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)

            # Process blockquotes
            elif tag_name == 'blockquote':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                p.style = 'Quote'
                p.add_run(element.get_text())

            # Process tables
            elif tag_name == 'table':
                # Count rows and columns
                rows = element.find_all('tr')
                if not rows:
                    continue

                # Determine max columns by checking all rows
                max_cols = 0
                for row in rows:
                    cols = row.find_all(['td', 'th'])
                    max_cols = max(max_cols, len(cols))

                if max_cols == 0:
                    continue

                # Create table
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'

                # Fill table
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j < max_cols:
                            # Style header cells
                            if cell.name == 'th':
                                cell_text = cell.get_text().strip()
                                table.cell(i, j).text = cell_text
                                for paragraph in table.cell(i, j).paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                            else:
                                table.cell(i, j).text = cell.get_text().strip()

        # Save document to a BytesIO object
        from io import BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue()

    def _apply_style(self, doc: Document, style_name: str) -> Document:
        """Apply predefined style to document."""
        # Default style is already applied by default

        if style_name == 'academic':
            # Font settings
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal'].font.size = Pt(12)

            # Heading styles
            for i in range(1, 4):
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = 'Times New Roman'
                heading_style.font.bold = True
                heading_style.font.size = Pt(16 - (i - 1) * 2)  # H1: 16pt, H2: 14pt, H3: 12pt

            # Paragraph spacing
            doc.styles['Normal'].paragraph_format.line_spacing = 2.0  # Double spacing

        elif style_name == 'business':
            # Font settings
            doc.styles['Normal'].font.name = 'Calibri'
            doc.styles['Normal'].font.size = Pt(11)

            # Heading styles
            for i in range(1, 4):
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = 'Calibri'
                heading_style.font.bold = True
                if i == 1:
                    heading_style.font.size = Pt(16)
                    heading_style.font.color.rgb = RGBColor(0, 77, 113)  # Dark blue
                elif i == 2:
                    heading_style.font.size = Pt(14)
                    heading_style.font.color.rgb = RGBColor(0, 112, 155)  # Medium blue
                else:
                    heading_style.font.size = Pt(12)
                    heading_style.font.color.rgb = RGBColor(0, 130, 188)  # Light blue

        elif style_name == 'minimal':
            # Font settings
            doc.styles['Normal'].font.name = 'Arial'
            doc.styles['Normal'].font.size = Pt(10)

            # Heading styles
            for i in range(1, 4):
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = 'Arial'
                heading_style.font.bold = True
                heading_style.font.size = Pt(14 - (i - 1) * 2)  # H1: 14pt, H2: 12pt, H3: 10pt

            # Paragraph spacing
            doc.styles['Normal'].paragraph_format.space_after = Pt(6)

        return doc

    def _create_table_of_contents(self, doc: Document) -> None:
        """Create table of contents"""
        doc.add_heading('Table of Contents', level=1)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update table of contents."

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)
        r_element.append(fldChar4)

        doc.add_page_break()

    def _add_hyperlink(self, paragraph, url, text):
        """Add a hyperlink to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                              is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0000FF')
        rPr.append(c)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink